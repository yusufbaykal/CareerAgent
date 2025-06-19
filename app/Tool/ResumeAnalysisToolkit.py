from pathlib import Path
import json
from agno.tools import Toolkit
from agno.utils.log import logger
from datetime import datetime
from typing import Optional


class ResumeAnalysisToolkit(Toolkit):
    
    def __init__(self):
        super().__init__(name="resume_tools")
        self.resume_dir = Path("Jobs/Resumes")
        self.results_dir = Path("Jobs/Resume_Analysis")
        self.register(self.parse_resume)
        self.register(self.analyze_resume)
        self.register(self.save_analysis)
        
    def _check_dependencies(self) -> dict:
        dependencies = {
            'PyPDF2': False,
            'docx': False
        }
        
        try:
            import PyPDF2
            dependencies['PyPDF2'] = True
        except ImportError:
            pass
            
        try:
            import docx
            dependencies['docx'] = True
        except ImportError:
            pass
            
        return dependencies
        
    def parse_resume(self, file_path: str) -> str:
        try:
            logger.info(f"İşlenen dosya yolu: {file_path}")
            
            possible_paths = self._generate_possible_paths(file_path)
            
            found_path = self._find_existing_file(possible_paths)
            if not found_path:
                return self._create_error_message(possible_paths)
            
            logger.info(f"✅ Dosya bulundu: {found_path}")
            
            return self._process_file_by_type(found_path)
                
        except Exception as e:
            error_msg = f"❌ Genel dosya okuma hatası: {str(e)}"
            logger.error(error_msg)
            return error_msg
    
    def _generate_possible_paths(self, file_path: str) -> list:
        possible_paths = []
        
        possible_paths.append(Path(file_path))
        
        if file_path.startswith("Jobs/Resumes/"):
            clean_path = file_path[13:]
            possible_paths.append(self.resume_dir / clean_path)
            possible_paths.append(Path(clean_path))
        
        file_name = Path(file_path).name
        possible_paths.append(self.resume_dir / file_name)
        
        path_obj = Path(file_path)
        if not path_obj.is_absolute():
            possible_paths.append(self.resume_dir / path_obj)
        
        return possible_paths
    
    def _find_existing_file(self, possible_paths: list) -> Optional[Path]:
        for path in possible_paths:
            logger.info(f"Denenen yol: {path}")
            if path.exists():
                return path
        return None
    
    def _create_error_message(self, possible_paths: list) -> str:
        error_msg = f"❌ Dosya bulunamadı. Denenen yollar:\n"
        for path in possible_paths:
            error_msg += f"  - {path}\n"
        logger.error(error_msg)
        return error_msg
    
    def _process_file_by_type(self, file_path: Path) -> str:
        suffix = file_path.suffix.lower()
        
        processors = {
            '.pdf': self._process_pdf,
            '.doc': self._process_word,
            '.docx': self._process_word,
            '.txt': self._process_text
        }
        
        processor = processors.get(suffix)
        if processor:
            return processor(file_path)
        else:
            error_msg = f"❌ Desteklenmeyen dosya formatı: {suffix}. Desteklenen formatlar: {', '.join(processors.keys())}"
            logger.error(error_msg)
            return error_msg
    
    def _process_pdf(self, file_path: Path) -> str:
        dependencies = self._check_dependencies()
        
        if not dependencies['PyPDF2']:
            error_msg = "❌ PDF dosyalarını işlemek için PyPDF2 kütüphanesi gereklidir. 'pip install PyPDF2' komutu ile yükleyebilirsiniz."
            logger.error(error_msg)
            return error_msg
        
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():  # Boş sayfa kontrolü
                            text += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"Sayfa {page_num + 1} okunamadı: {str(e)}")
                        continue
            
            if not text.strip():
                return "❌ PDF dosyasından metin çıkarılamadı. Dosya görüntü tabanlı olabilir."
            
            logger.info(f"✅ PDF başarıyla okundu: {len(text)} karakter")
            return text
                
        except Exception as e:
            error_msg = f"❌ PDF okuma hatası: {str(e)}"
            logger.error(error_msg)
            return error_msg
    
    def _process_word(self, file_path: Path) -> str:
        dependencies = self._check_dependencies()
        
        if not dependencies['docx']:
            error_msg = "❌ Word dosyalarını işlemek için python-docx kütüphanesi gereklidir. 'pip install python-docx' komutu ile yükleyebilirsiniz."
            logger.error(error_msg)
            return error_msg
        
        try:
            import docx
            
            doc = docx.Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            if not text.strip():
                return "❌ Word belgesinden metin çıkarılamadı."
            
            logger.info(f"✅ Word belgesi başarıyla okundu: {len(text)} karakter")
            return text
                
        except Exception as e:
            error_msg = f"❌ Word belgesi okuma hatası: {str(e)}"
            logger.error(error_msg)
            return error_msg
    
    def _process_text(self, file_path: Path) -> str:
        try:
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    text = file.read()
            except UnicodeDecodeError:
                logger.warning("UTF-8 okuma başarısız, latin-1 deneniyor")
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
            
            if not text.strip():
                return "❌ Metin dosyası boş."
            
            logger.info(f"✅ Metin dosyası başarıyla okundu: {len(text)} karakter")
            return text
                
        except Exception as e:
            error_msg = f"❌ Metin dosyası okuma hatası: {str(e)}"
            logger.error(error_msg)
            return error_msg
    
    def analyze_resume(self, resume_content: str) -> str:
        return resume_content
        
    def save_analysis(self, analysis_data: str, resume_name: str) -> str:
        try:
            self.results_dir.mkdir(exist_ok=True, parents=True)
            
            file_name = self._generate_filename(resume_name, analysis_data)
            file_path = self.results_dir / file_name
            
            data = self._prepare_analysis_data(analysis_data, resume_name)
                
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
                
            logger.info(f"✅ Analiz sonucu kaydedildi: {file_path}")
            return f"✅ Analiz başarıyla {file_path} dosyasına kaydedildi."
            
        except Exception as e:
            error_msg = f"❌ Analiz kaydedilirken hata oluştu: {str(e)}"
            logger.error(error_msg)
            return error_msg
    
    def _generate_filename(self, resume_name: str, analysis_data: str) -> str:
        if resume_name.endswith('.json') and '_' in resume_name:
            return resume_name
        else:
            base_name = Path(resume_name).stem
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            return f"resume_analysis_{base_name}_{timestamp}.json"
    
    def _prepare_analysis_data(self, analysis_data: str, resume_name: str) -> dict:
        if isinstance(analysis_data, str):
            try:
                return json.loads(analysis_data)
            except json.JSONDecodeError:
                return {
                    "status": "success",
                    "analysis": analysis_data,
                    "timestamp": datetime.now().isoformat(),
                    "file_name": resume_name,
                    "metadata": {
                        "processed_by": "ResumeAnalysisToolkit",
                        "version": "1.0"
                    }
                }
        else:
            return analysis_data if isinstance(analysis_data, dict) else {}